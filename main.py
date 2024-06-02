from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import logging
import os

class GeradorCertificado:
    """
    Construtor da classe
    :param modelo_certificado = Modelo do certificado usado para geração de certificados em massa
    :param planilha_alunos = Planilha do excel que contém os nomes dos alunos que receberam os certificados
    :param sheet_name = Nome da planilha dos alunos
    """
    def __init__(self, modelo_certificado="Certificado1.docx", planilha_alunos="Alunos.xlsx", sheet_name="Nomes"):
        self.modelo_certificado = modelo_certificado
        self.planilha_alunos = planilha_alunos

        try:
            if not os.path.exists(modelo_certificado):
                raise FileNotFoundError(f"Modelo de certificado {modelo_certificado} não encontrado.")
            if not os.path.exists(planilha_alunos):
                raise FileNotFoundError(f"Planilha de alunos {planilha_alunos} não encontrada.")

            self.arquivo_certificado = Document(modelo_certificado)
            self.arquivo_alunos = pd.read_excel(planilha_alunos, sheet_name=sheet_name)
            self.estilo = self.arquivo_certificado.styles["Normal"]
        except FileExistsError as e:
            logging.error("Arquivo não encontrado: %s", e)
        except Exception as e:
            logging.error("Erro na inicialização dos certificados: %s", e)

    def _personalizar_certificado(self, documento, nome):
        """
        Personaliza o certificado substituindo o placeholder pelo nome do aluno.
        :param documento: Documento do Word a ser personalizado.
        :param nome: Nome do aluno para inserir no certificado.
        """
        for paragrafo in documento.paragraphs:
            if "@nome" in paragrafo.text:
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for fonte in paragrafo.runs:
                    fonte.text = fonte.text.replace("@nome", nome)
                    fonte.font.name = "Calibri (Corpo)"
                    fonte.font.size = Pt(24)
                    fonte.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                    fonte.font.bold = True
    def gerando_certificado(self):
        """
        Método gerador de certificados
        """
        if self.arquivo_alunos is not None:
            try:
                for index, row in self.arquivo_alunos.iterrows():
                    nome = row["Aluno"]
                    if nome:
                        novo_arquivo = Document(self.modelo_certificado)
                        self._personalizar_certificado(novo_arquivo, nome)
                        novo_arquivo.save(f"Certificado_{nome}.docx")
                    else:
                        logging.warning("Linha %d na planilha não contém nome do aluno.", index)
                logging.info("Certificados gerados com sucesso!")
            except KeyError as e:
                logging.error("Coluna 'Aluno' não encontrada na planilha: %s", e)
            except Exception as e:
                logging.error("Ocorreu um erro: %s", e)


if __name__ == "__main__":

    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
    try:
        res = GeradorCertificado()
        res.gerando_certificado()
    except Exception as e:
        logging.critical("Erro crítico: %s", e)
